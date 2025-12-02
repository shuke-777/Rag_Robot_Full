from langchain_openai import ChatOpenAI
from dotenv import load_dotenv
load_dotenv()
import os
from pathlib import Path
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document as langchainDocument
from langchain_classic.chains.history_aware_retriever import create_history_aware_retriever
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.messages import HumanMessage, AIMessage
from langchain_classic.chains.combine_documents import create_stuff_documents_chain
from langchain_core.runnables import RunnableParallel
from langchain_community.chat_message_histories import ChatMessageHistory
from langchain_core.runnables import RunnableWithMessageHistory
from langchain_community.embeddings.zhipuai import ZhipuAIEmbeddings
import asyncio

# 项目根目录
PROJECT_ROOT = Path(__file__).parent.parent
DATA_DIR = PROJECT_ROOT / "data"
VECTOR_STORE_DIR = PROJECT_ROOT / "vector_store"

# 初始化LLM
llm = ChatOpenAI(
    model_name=os.getenv('ZHIPU_MODEL_NAME'),
    api_key=os.getenv('ZHIPUAI_API_KEY'),
    base_url=os.getenv('ZHIPUAI_URL')
)

# 初始化Embedding模型
embed_model = ZhipuAIEmbeddings(
    model_name='embedding-2',
    api_key=os.getenv('ZHIPUAI_API_KEY')
)


def load_split_documents(documents_name: str) -> list:
    """加载并分割Word文档"""
    read_text = Document(documents_name)
    result = [para.text for para in read_text.paragraphs if para.text.strip()]
    full_result = [langchainDocument(page_content=text) for text in result]
    text_split = RecursiveCharacterTextSplitter(chunk_size=200, chunk_overlap=20)
    result1 = text_split.split_documents(full_result)
    return result1


def create_vector_store(documents: list, store_name: str = "faiss_index"):
    """创建并保存向量存储"""
    embed = FAISS.from_documents(documents, embed_model)
    save_path = VECTOR_STORE_DIR / store_name
    embed.save_local(str(save_path))
    return embed


def load_vector_store(store_name: str = "faiss_index"):
    """加载已存在的向量存储"""
    load_path = VECTOR_STORE_DIR / store_name
    embed = FAISS.load_local(
        str(load_path),
        embed_model,
        allow_dangerous_deserialization=True
    )
    return embed


def create_history_aware_retriever_chain(llm, retriever):
    """创建历史感知检索器"""
    prompt_text = """
    你是一个问题重构助手，现在你根据上下文来进行问题重构，
    要求:
    1.重构问题
    2.不要回答问题
    3.保持问题的严谨性
    """
    prompt = ChatPromptTemplate.from_messages([
        ('system', prompt_text),
        MessagesPlaceholder(variable_name='chat_history'),
        ('human', '{input}')
    ])

    history_aware_retriever = create_history_aware_retriever(
        llm=llm,
        retriever=retriever,
        prompt=prompt
    )
    return history_aware_retriever


def create_qa_chain(llm, history_aware_retriever):
    """创建问答链"""
    prompt_text = """
    你是一个无所不能的助手，根据我的上下文来进行回答{context}。如果不知道，请说不知道，直接告诉我答案就好了
    """
    qa_prompt = ChatPromptTemplate.from_messages([
        ('system', prompt_text),
        MessagesPlaceholder(variable_name='chat_history'),
        ('human', '{input}')
    ])

    doc_chain = create_stuff_documents_chain(llm=llm, prompt=qa_prompt)
    qa_full_chain = (
        RunnableParallel({
            'context': history_aware_retriever,
            'chat_history': lambda x: x['chat_history'],
            'input': lambda x: x['input']
        }) | doc_chain
    )
    return qa_full_chain

def invoke_limit_history(chain,question, chat_history,max_rounds,):
    """
    手动限制聊天记录长度
    """
    max_messages = max_rounds * 2
    if len(chat_history)>max_messages:
        history = chat_history[-max_messages:]
    else:
        history = chat_history
    limit_result_history = chain.invoke({'input':question,'chat_history':history})
    return limit_result_history

# 会话存储
store = {}

##在原有的基础上，限制历史长度
def get_session_history(session_id: str):
    """获取或创建会话历史"""
    try:
        if session_id not in store:
            store[session_id] = ChatMessageHistory()
        history = store[session_id]
        if len(history.messages)>20:
            #保留最近20条记录
            history.messages = history.messages[-20:]
        return history
    except Exception as e:
        print(f"Error: {e}")
        return ChatMessageHistory()

def get_rag_chain(rag_chain):
    """创建带历史管理的RAG链"""
    rag_chain_with_history = RunnableWithMessageHistory(
        rag_chain,
        get_session_history,
        input_messages_key='input',
        history_messages_key='chat_history'
    )
    return rag_chain_with_history

#stream/async
async def astream_invoke(rag_chain,question: list,max_rounds:int,chat_history):
    """异步流式调用"""
    max_messages = max_rounds * 2
    if len(chat_history)>max_messages:
        history = chat_history[-max_messages:]
    else:
        history = chat_history
    for content in question:
        yield f'\n问题:{content}\n回答:'
        async for chunk in rag_chain.astream({'input':content,'chat_history':history}):
            yield chunk
#手动管理history的func
async def astream_invoke_main(rag_chain,question: list,max_rounds:int,chat_history):
    async for chunk in astream_invoke(rag_chain,question,max_rounds,chat_history):
        print(chunk,end = '',flush=True)

async def chat_with_input_astream(rag_chain,session_id):
    print('聊天系统启动')
    while True:
        try:
            input_question = input('用户提问:\n')

            if input_question.lower() == 'quit':
                print('再见')
                return
           # yield f'用户提问:{input_question}\n'
            yield f'AI回答:\n'
            input_rag_chain=get_rag_chain(rag_chain)
            async for chunk in input_rag_chain.astream({'input':input_question},
                                                       config ={'configurable':{'session_id':session_id}}):
                yield chunk
            yield '\n'
        except Exception as e:
            print(f"系统出现错误: {e}")
#用session来管理聊天记录
async def astream_session_invoke_main(rag_chain,session_id):
    async for chunk in chat_with_input_astream(rag_chain,session_id):
        print(chunk,end = '',flush=True)
    print(get_session_history(session_id))

if __name__ == '__main__':
    # 示例数据文件路径
    doc_path = DATA_DIR / "解析木--实验计划.docx"

    print('-' * 30, '1_手动管理聊天记录')

    # 加载文档（首次运行时需要创建向量存储）
    t1_split_text = load_split_documents(str(doc_path))

    # 创建向量存储（首次运行时取消注释）
    # create_vector_store(t1_split_text, "1201Faiss.faiss")

    # 加载已有的向量存储
    embed = load_vector_store("1201Faiss.faiss")
    retrieval = embed.as_retriever(search_kwargs={'k': 3})

    # 创建检索和问答链
    aware_history = create_history_aware_retriever_chain(llm, retrieval)
    qa_chain = create_qa_chain(llm, aware_history)

    # 手动管理聊天历史
    chat_history = []
    question1 = '解析木怎么选择'
    #result1 = qa_chain.invoke({'input': question1, 'chat_history': chat_history})
    # chat_history.extend([
    #     HumanMessage(content=question1),
    #     AIMessage(content=result1)
    # ])

    question2 = '它的密度单位是什么'
    # result2 = qa_chain.invoke({'input': question2, 'chat_history': chat_history})
    # chat_history.extend([
    #     HumanMessage(content=question2),
    #     AIMessage(content=result2)
    # ])

    print('-' * 30, '2_自动管理聊天记录')
    session_id = 'user_1'
    # rag_chain1 = get_rag_chain(qa_chain)
    #
    # rag_chain1.invoke(
    #     {'input': question1},
    #     config={'configurable': {'session_id': session_id}}
    # )
    # rag_chain1.invoke(
    #     {'input': question2},
    #     config={'configurable': {'session_id': session_id}}
    # )

    #print('查看历史记录')
    history = get_session_history(session_id)
    #print(history)
    for i, message in enumerate(history.messages):
        msg_type = '用户' if isinstance(message, HumanMessage) else 'AI'
        print(f'{i},[{msg_type}]:{message.content}', end='\n\n')

    print('-' * 30, '3_限制历史记录长度:')
    # 手动管理聊天历史
    # chat_history1 = []
    # question1 = '解析木怎么选择'
    # result1 = invoke_limit_history(qa_chain, question1, chat_history1, 10)
    # chat_history1.extend([
    #     HumanMessage(content=question1),
    #     AIMessage(content=result1)
    # ])

    # question2 = '它的密度单位是什么'
    # #result1 = invoke_limit_history(qa_chain, question1, chat_history1, 10)
    # chat_history1.extend([
    #     HumanMessage(content=question2),
    #     AIMessage(content=result2)
    # ])
    #print(chat_history1)
    print('-' * 30, '4_stream/async')
    question_list = [question1,question2]
    #asyncio.run(astream_invoke_main(qa_chain, question_list, max_rounds=10,chat_history=chat_history1))
    print('-'*30,'5_session增加对话处理，增加input')
    asyncio.run(astream_session_invoke_main(qa_chain, session_id))
